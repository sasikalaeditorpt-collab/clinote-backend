from pathlib import Path
from docx import Document

def build_styled_docx(styled_text: str, template_path: Path | None, output_path: Path) -> Path:
    if template_path and template_path.exists():
        doc = Document(str(template_path))
        # Optional: clear body if template has placeholder text
        for p in doc.paragraphs:
            if p.text.strip():
                p.text = ""
    else:
        doc = Document()

    for line in styled_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")

    doc.save(str(output_path))
    return output_path